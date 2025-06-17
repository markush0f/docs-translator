import requests
import uuid
import json
import os
from dotenv import load_dotenv
from docx import Document

# Cargar las variables de entorno
load_dotenv()

# Configuración de Azure Translator
key = os.getenv("AZURE_TRANSLATOR_KEY")
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "northeurope"

path = "/translate"
constructed_url = endpoint + path

params = {
    "api-version": "3.0",
    "from": "en",  # Idioma de origen
    "to": "es",  # Idioma de destino
}

headers = {
    "Ocp-Apim-Subscription-Key": key,
    "Ocp-Apim-Subscription-Region": location,
    "Content-type": "application/json",
    "X-ClientTraceId": str(uuid.uuid4()),
}


# Función para enviar texto a la API de Azure para traducción
def translate_text(text: str, target_lang: str) -> str:
    body = [{"text": text}]
    params["to"] = target_lang  # Establecer el idioma de destino dinámicamente
    response = requests.post(constructed_url, params=params, headers=headers, json=body)
    response_json = response.json()

    # Retornar el texto traducido
    return response_json[0]["translations"][0]["text"]


# Función para traducir un archivo DOCX
def translate_docx(input_path: str, output_path: str, target_lang: str):
    doc = Document(input_path)

    # Concatenar todo el texto del documento
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"

    # Traducir todo el texto del documento
    translated_text = translate_text(full_text, target_lang)

    # Insertar el texto traducido de vuelta en el documento
    translated_paragraphs = translated_text.split("\n")
    para_index = 0
    for para in doc.paragraphs:
        if para_index < len(translated_paragraphs):
            para.text = translated_paragraphs[para_index]
            para_index += 1

    # Guardar el documento traducido
    doc.save(output_path)

    print(f"Documento traducido guardado como: {output_path}")
