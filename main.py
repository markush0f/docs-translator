import requests
import uuid
import json
import os
from dotenv import load_dotenv
from docx import Document
from docx.oxml.ns import qn

load_dotenv()

key = os.getenv("AZURE_TRANSLATOR_KEY")
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "northeurope"

path = "/translate"
constructed_url = endpoint + path

params = {
    "api-version": "3.0",
    "from": "es",
    "to": "en",
}

headers = {
    "Ocp-Apim-Subscription-Key": key,
    "Ocp-Apim-Subscription-Region": location,
    "Content-type": "application/json",
    "X-ClientTraceId": str(uuid.uuid4()),
}


# Función para enviar texto a la API de Azure para traducción
def translate_text(text: str, target_lang: str) -> str:
    body = [{"text": text}]
    params["to"] = target_lang  # Establecer el idioma de destino dinámicamente
    response = requests.post(constructed_url, params=params, headers=headers, json=body)
    response_json = response.json()

    # Retornar el texto traducido
    return response_json[0]["translations"][0]["text"]


# Función para comprobar si un "run" tiene una imagen
def has_drawing(run) -> bool:
    return bool(run._element.findall(".//" + qn("w:drawing")))


def translate_paragraph(paragraph, target_lang: str):
    full_text = paragraph.text
    if not full_text.strip():
        return

    translated_text = translate_text(full_text, target_lang)

    # Reemplazar el texto traducido en el párrafo, manteniendo el formato
    written = False
    for run in paragraph.runs:
        if run.text.strip() and not has_drawing(run):
            if not written:
                run.text = translated_text
                written = True
            else:
                run.text = ""


def translate_paragraphs(paragraphs, target_lang: str):
    for para in paragraphs:
        translate_paragraph(para, target_lang)


def translate_tables(tables, target_lang: str):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                translate_paragraphs(cell.paragraphs, target_lang)


def translate_section_parts(section, target_lang: str):
    parts = [
        section.header,
        section.footer,
        section.first_page_header,
        section.first_page_footer,
        section.even_page_header,
        section.even_page_footer,
    ]
    for part in parts:
        translate_paragraphs(part.paragraphs, target_lang)
        translate_tables(part.tables, target_lang)


# Función para traducir todo el documento
def translate_docx(input_path: str, output_path: str, target_lang: str):
    doc = Document(input_path)

    # Traducir todos los párrafos del cuerpo principal del documento
    translate_paragraphs(doc.paragraphs, target_lang)

    # Traducir las tablas
    translate_tables(doc.tables, target_lang)

    # Traducir encabezados y pies de página de todas las secciones
    for section in doc.sections:
        translate_section_parts(section, target_lang)

    doc.save(output_path)


input_path = "document.docx"
output_path = "translated_document.docx"
target_lang = "en"

translate_docx(input_path, output_path, target_lang)
print(f"El documento traducido se ha guardado en: {output_path}")
