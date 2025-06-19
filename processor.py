from docx import Document
from docx.oxml.ns import qn

def has_drawing(run) -> bool:
    return bool(run._element.findall(".//" + qn("w:drawing")))

def translate_paragraph(paragraph, target_lang, translate_func):
    full_text = paragraph.text
    if not full_text.strip():
        return

    translated_text = translate_func(full_text, target_lang)
    written = False
    for run in paragraph.runs:
        if run.text.strip() and not has_drawing(run):
            if not written:
                run.text = translated_text
                written = True
            else:
                run.text = ""

def translate_paragraphs(paragraphs, target_lang, translate_func):
    for para in paragraphs:
        translate_paragraph(para, target_lang, translate_func)

def translate_tables(tables, target_lang, translate_func):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                translate_paragraphs(cell.paragraphs, target_lang, translate_func)

def translate_section_parts(section, target_lang, translate_func):
    parts = [
        section.header,
        section.footer,
        section.first_page_header,
        section.first_page_footer,
        section.even_page_header,
        section.even_page_footer,
    ]
    for part in parts:
        translate_paragraphs(part.paragraphs, target_lang, translate_func)
        translate_tables(part.tables, target_lang, translate_func)

def translate_docx(input_path: str, output_path: str, target_lang: str, translate_func):
    doc = Document(input_path)
    translate_paragraphs(doc.paragraphs, target_lang, translate_func)
    translate_tables(doc.tables, target_lang, translate_func)
    for section in doc.sections:
        translate_section_parts(section, target_lang, translate_func)
    doc.save(output_path)
