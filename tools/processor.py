from config.logger_config import setup_logger
from docx import Document
from docx.oxml.ns import qn

logger = setup_logger("Processor")


def has_drawing(run) -> bool:
    return bool(run._element.findall(".//" + qn("w:drawing")))


def translate_paragraph(paragraph, target_lang, source_lang, translate_func):
    full_text = paragraph.text
    if not full_text.strip():
        logger.debug("Párrafo vacío, se omite.")
        return

    logger.debug(f"Traduciendo párrafo: {full_text[:40]}...")
    translated_text = translate_func(full_text, target_lang, source_lang)
    written = False
    for run in paragraph.runs:
        if run.text.strip() and not has_drawing(run):
            if not written:
                run.text = translated_text
                written = True
            else:
                run.text = ""


def translate_paragraphs(paragraphs, target_lang, source_lang, translate_func):
    logger.debug(f"Traduciendo {len(paragraphs)} párrafos.")
    for para in paragraphs:
        translate_paragraph(para, target_lang, source_lang, translate_func)


def translate_tables(tables, target_lang, source_lang, translate_func):
    logger.debug(f"Traduciendo {len(tables)} tablas.")
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                translate_paragraphs(
                    cell.paragraphs, target_lang, source_lang, translate_func
                )


def translate_section_parts(section, target_lang, source_lang, translate_func):
    logger.debug("Traduciendo secciones: encabezados y pies de página.")
    parts = [
        section.header,
        section.footer,
        section.first_page_header,
        section.first_page_footer,
        section.even_page_header,
        section.even_page_footer,
    ]
    for part in parts:
        translate_paragraphs(part.paragraphs, target_lang, source_lang, translate_func)
        translate_tables(part.tables, target_lang, source_lang, translate_func)


def translate_docx(
    input_path: str,
    output_path: str,
    target_lang: str,
    source_lang: str,
    translate_func,
):
    logger.info(f"Iniciando traducción de: {input_path}")
    doc = Document(input_path)

    translate_paragraphs(doc.paragraphs, target_lang, source_lang, translate_func)
    translate_tables(doc.tables, target_lang, source_lang, translate_func)

    for section in doc.sections:
        translate_section_parts(section, target_lang, source_lang, translate_func)

    doc.save(output_path)
    logger.info(f"Documento traducido guardado en: {output_path}")
