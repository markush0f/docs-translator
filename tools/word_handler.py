import requests
import uuid
import json
import os
from dotenv import load_dotenv
from docx import Document

from types_doc import ParamsType


# Cargar las variables de entorno
load_dotenv()

# Configuración de Azure Translator
key = os.getenv("AZURE_TRANSLATOR_KEY")
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "northeurope"

path = "/translate"
constructed_url = endpoint + path


headers = {
    "Ocp-Apim-Subscription-Key": key,
    "Ocp-Apim-Subscription-Region": location,
    "Content-type": "application/json",
    "X-ClientTraceId": str(uuid.uuid4()),
}


def translate_text(text: str, target_lang: str, source_lang: str) -> str:
    params = {
        "api-version": "3.0",
        "from": source_lang,
        "to": target_lang,
    }
    body = [{"text": text}]
    response = requests.post(constructed_url, params=params, headers=headers, json=body)
    response_json = response.json()

    return response_json[0]["translations"][0]["text"]


# Función para traducir un archivo DOCX
def translate_docx(input_path: str, output_path: str, target_lang: str) -> None:
    doc = Document(input_path)
    source_lang = "en"
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"

    translated_text = translate_text(full_text, target_lang, source_lang)

    translated_paragraphs = translated_text.split("\n")
    para_index = 0
    for para in doc.paragraphs:
        if para_index < len(translated_paragraphs):
            para.text = translated_paragraphs[para_index]
            para_index += 1

    doc.save(output_path)

    print(f"Documento traducido guardado como: {output_path}")
